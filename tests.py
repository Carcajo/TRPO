import pytest
from unittest.mock import AsyncMock, MagicMock, patch
from aiogram.types import Message, Document, File, ContentType
from io import BytesIO
from docx import Document
from main import (
    calculate_similarity,
    extract_text_from_docx,
    is_text_generated_by_agi,
    request_text,
    analyze_text,
    send_welcome,
)


@pytest.mark.asyncio
async def test_send_welcome():
    message = AsyncMock()
    message.chat.id = 12345
    state = AsyncMock()

    await send_welcome(message, state)

    message.reply.assert_called_with(
        "Привет, я помогу проверить текст. Выбери опцию.",
        reply_markup=pytest.ANY
    )
    state.set_state.assert_called_with("BotStates:MAIN_MENU")


def test_calculate_similarity():
    original = "Это тестовый текст для проверки"
    generated = "Это текст для теста проверки"
    result = calculate_similarity(original, generated)
    assert result == 4 / 5


def test_extract_text_from_docx():
    docx_file = BytesIO()
    doc = Document()
    doc.add_paragraph("Первый абзац")
    doc.add_paragraph("Второй абзац")
    doc.save(docx_file)
    docx_file.seek(0)

    result = extract_text_from_docx(docx_file)

    assert result == "Первый абзац\nВторой абзац\n"


@pytest.mark.asyncio
async def test_request_text():
    message = AsyncMock()
    message.chat.id = 12345
    state = AsyncMock()

    await request_text(message, state)

    message.reply.assert_called_with(
        "Пришли текст или файл и я проверю его.",
        reply_markup=None
    )
    state.set_state.assert_called_with("BotStates:WAITING_FOR_TEXT")


@pytest.mark.asyncio
@patch("main.bot.get_file")
@patch("main.bot.download_file")
async def test_analyze_text(mock_download_file, mock_get_file):
    message = AsyncMock()
    message.content_type = ContentType.TEXT
    message.text = "Тестовый текст"
    state = AsyncMock()

    await analyze_text(message, state)

    message.reply.assert_called_with("Идет анализ...")
    message.reply.assert_any_call(pytest.ANY)


@pytest.mark.asyncio
@patch("main.bot.get_file")
@patch("main.bot.download_file")
async def test_analyze_docx_file(mock_download_file, mock_get_file):
    message = AsyncMock()
    message.content_type = ContentType.DOCUMENT
    message.document = MagicMock()
    message.document.file_name = "test.docx"
    state = AsyncMock()

    mock_file_content = BytesIO()
    doc = Document()
    doc.add_paragraph("Пример текста в файле")
    doc.save(mock_file_content)
    mock_file_content.seek(0)
    mock_download_file.return_value = mock_file_content

    await analyze_text(message, state)

    message.reply.assert_any_call(pytest.ANY)
